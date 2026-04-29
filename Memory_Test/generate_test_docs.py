#!/usr/bin/env python3
"""
Generate 100 fictional company-internal documents for Memori-Vault retrieval testing.

Purpose:
- Avoid public/common-knowledge articles.
- Make answers verifiable only from indexed local documents.
- Cover md/txt/docx/pdf ingestion and sidebar display.

All names, projects, people, dates, indicators and policies below are fictional test data.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors

OUTPUT_DIR = Path(__file__).resolve().parent
COMPANY = "星衡智能"
DISCLAIMER = "【虚构内部测试资料】本文件仅用于 Memori-Vault 本地检索评测，不包含真实公司机密。"


@dataclass(frozen=True)
class Dossier:
    category: str
    project: str
    code: str
    owner: str
    department: str
    date: str
    anchor_fact: str
    metric: str
    workflow: str
    exception: str
    wrong_answer_guard: str


DOSSIERS = [
    Dossier("产品策略", "极光账本", "AUR-17", "林知远", "企业产品部", "2026-01-08", "极光账本的试点客户只允许选择 3 家：禾川物流、南屿商贸、钧石制造。", "北极星指标不是 DAU，而是“月度已核销对账单数”，一期目标为 18,600 张。", "需求冻结后，任何字段改名必须先通过“字段契约评审”，再进入灰度白名单。", "禾川物流拥有一次“免评审字段别名”例外，但只能用于 customer_alias 字段。", "不要按通用 SaaS 项目写成 DAU 或转化率。"),
    Dossier("运营增长", "青梧留存", "GROW-29", "陈夏", "增长运营组", "2026-01-12", "公司内部把“冷启动用户”定义为 14 天内访问过但未创建项目的人，不是刚注册用户。", "青梧留存的核心阈值是 D14 二次创建率 31.5%，低于 29% 必须暂停投放。", "每周二 10:30 先看 cohort，再决定素材预算；不得先看点击率。", "教育行业线允许将暂停阈值放宽到 27.8%，但必须由陈夏签字。", "不要回答行业常见的新注册用户定义。"),
    Dossier("安全合规", "玄武密钥", "SEC-41", "周曼青", "安全办公室", "2026-01-16", "玄武密钥轮换窗口固定为每月第二个周四 22:10-22:35。", "服务端密钥最大存活期为 43 天；移动端派生密钥为 9 天。", "轮换前 6 小时必须完成影子校验，校验记录写入 xh_key_shadow_audit。", "春节冻结期只允许轮换 payment-salt，不允许轮换 tenant-root。", "不要写成常见的 90 天轮换。"),
    Dossier("技术架构", "银杏-17", "ARC-17", "苏澈", "平台架构组", "2026-01-19", "银杏-17 的上线冻结窗口是每周三 19:40-21:10。", "网关错误率超过 0.47% 且持续 6 分钟才触发自动回滚。", "发布前必须先执行 shadow-read，再执行 2.7% 灰度，最后才到 12% 灰度。", "如果依赖开票服务 invoice-core，则禁止在周三窗口上线。", "不要回答成常见的 5% 灰度或任意时间发布。"),
    Dossier("客户成功", "白鹭工单", "CS-08", "马溪", "客户成功部", "2026-01-23", "白鹭工单的 P1 升级 SLA 是 11 分钟，P2 是 37 分钟。", "一线首次响应满意度目标为 92.4%，但不纳入月奖，只纳入复盘。", "P1 先电话确认再建群；P2 先补齐租户号再同步研发值班。", "鲸舟银行租户的 P2 可按 P1 通知马溪，但 SLA 仍按 37 分钟统计。", "不要套用通用 15/30/60 分钟工单 SLA。"),
    Dossier("财务风控", "赤松预算", "FIN-33", "何砚", "财务风控组", "2026-01-26", "赤松预算的审批线不是金额固定线，而是“预算池剩余率”低于 18% 才升级 CFO。", "试点期单笔采购上限 78,000 元；但云资源预留实例按季度池计算。", "申请人先填写成本标签，再由系统匹配预算池，不允许先找财务口头确认。", "研发应急热修可先占用 RND-EMG 池，但 24 小时内必须补单。", "不要回答成传统金额越大越高层审批。"),
    Dossier("数据治理", "松塔标签", "DATA-12", "顾安", "数据平台组", "2026-02-02", "松塔标签把“活跃门店”定义为 10 天内完成过库存盘点的门店。", "标签产出延迟 SLO 是 T+38 分钟，不是 T+1 天。", "标签上线必须同时登记口径、回填范围、下游看板和撤销策略四项。", "华南试验区允许活跃门店口径临时放宽到 13 天，截止 2026-03-31。", "不要按零售行业常见日活门店口径回答。"),
    Dossier("人力组织", "石竹梯队", "HR-22", "赵闻笛", "组织发展部", "2026-02-06", "石竹梯队的 L3 晋升必要条件是主持过一次跨部门复盘，而不是带人数量。", "2026 上半年 L3 名额上限为 14 人，技术序列最多占 6 人。", "候选人材料必须包含“失败案例复盘页”，缺失则自动延期一轮。", "安全办公室可额外推荐 1 名候选人，但仍占总名额。", "不要写成通用绩效评分或团队规模决定晋升。"),
    Dossier("供应链", "蓝鲸 B17", "SCM-17", "唐稚", "供应链协同组", "2026-02-10", "蓝鲸 B17 的安全库存不是 30 天，而是 143 台关键模组。", "当 Bin-7 库位低于 51 台时触发黄色预警，低于 23 台触发红色预警。", "红色预警先通知唐稚，再通知采购，不允许绕过供应链协同组。", "若供应商为“柏屿电子”，黄色预警阈值提高到 67 台。", "不要按制造业常见天数库存回答。"),
    Dossier("品牌市场", "雾凇发布", "MKT-05", "沈乔", "品牌市场部", "2026-02-14", "雾凇发布的禁用词清单包含“颠覆行业”“零成本替代”“全自动合规”。", "发布会目标不是线索数，而是 46 家目标客户中至少 19 家完成闭门演示预约。", "对外物料先过法务术语检查，再过品牌语气检查，顺序不能反。", "政企版海报可保留“可信部署”四字，但必须删除“零风险”。", "不要输出营销通用夸张词。"),
    Dossier("法务合同", "琥珀条款", "LAW-19", "魏蓝", "法务部", "2026-02-18", "琥珀条款要求所有试点合同加入“本地模型不出域”附件 A-3。", "合同回传时限是 5 个工作日；超过后自动标记为 amber-delay。", "销售不得口头承诺模型效果，只能引用附件 A-3 的可验证边界。", "存量客户续签可暂不附 A-3，但新增模块必须附。", "不要写成普通 SaaS 数据处理协议即可。"),
    Dossier("研发效能", "蜂巢流水线", "DEV-26", "于北辰", "研发效能组", "2026-02-21", "蜂巢流水线的强制阻断项是“迁移脚本缺少回滚段”，不是单测覆盖率。", "主干构建 18 分钟内完成算绿色；超过 24 分钟算黄色；超过 31 分钟算红色。", "每次红色构建必须在 hcx_build_review 表登记根因标签。", "文档仓库允许黄色构建合并，但产品仓库不允许。", "不要只回答提高单测覆盖率。"),
    Dossier("售前方案", "云帆模板", "PRE-09", "梁其", "解决方案部", "2026-02-25", "云帆模板中“私有化部署”默认写作 3 节点轻量集群，不写 5 节点。", "标准演示时长为 28 分钟，其中知识库问答必须控制在 7 分钟内。", "售前不得展示未开启证据面板的回答截图。", "金融客户可延长到 36 分钟，但必须保留 7 分钟证据讲解。", "不要按通用售前方案写成越完整越好。"),
    Dossier("质量保障", "晴川验收", "QA-31", "董霁", "质量保障部", "2026-03-01", "晴川验收把“可交付”定义为完成 9 条端到端证据链，而不是测试用例全部通过。", "验收抽样比例固定为 17%，但核心流程必须 100% 覆盖。", "缺陷关闭前必须附带复现实录、修复提交和二次验证三项。", "UI 文案缺陷可以合并关闭，但检索证据缺陷不能合并关闭。", "不要回答成普通测试通过率。"),
    Dossier("内控审计", "镜湖审计", "AUD-14", "郝映", "内控审计组", "2026-03-05", "镜湖审计的抽查样本来自“最近 5 次异常回滚”，不是随机工单。", "每次审计至少追踪 4 个角色：申请人、审批人、执行人、复核人。", "审计结论必须写明“系统证据”和“人工解释”的分界。", "若异常回滚少于 5 次，则补足最近 30 天的灰度暂停记录。", "不要套用通用随机抽样审计。"),
    Dossier("知识管理", "栈桥知识库", "KM-07", "黎鸥", "知识管理组", "2026-03-09", "栈桥知识库的过期规则是 75 天无人引用即进入复审，不是按创建时间一年。", "核心知识卡必须包含适用场景、反例和负责组三项。", "复审先由知识所有人判断，再由使用频次最高的团队确认。", "安全类知识卡引用次数再低也不自动过期，只进入人工复审。", "不要回答普通文档生命周期管理。"),
    Dossier("渠道生态", "翠微伙伴", "CHN-11", "秦若", "渠道生态部", "2026-03-13", "翠微伙伴的分级依据是“联合交付成功次数”，不是签约金额。", "银级需要 2 次成功交付，金级需要 5 次，黑松级需要 9 次。", "伙伴提交案例后先由交付经理确认，再由渠道生态部定级。", "高校实验室伙伴可跳过银级，但不能跳过金级。", "不要按销售额或返点比例回答。"),
    Dossier("隐私评估", "芦苇脱敏", "PRI-28", "薛宁", "隐私委员会", "2026-03-17", "芦苇脱敏把门店联系人电话保留后 2 位用于排障，其余全部掩码。", "脱敏任务超过 13 分钟未完成要进入人工复核队列。", "任何导出样本必须先打上 sample_scope_id，再进入审批流。", "灰度排障期间可临时保留城市字段，但不得保留详细地址。", "不要写成全部字段简单删除。"),
    Dossier("应急响应", "灯塔演练", "INC-44", "任栎", "SRE 值班组", "2026-03-21", "灯塔演练规定主责人在 4 分钟内声明 incident commander。", "SEV-2 的内部同步频率是每 17 分钟一次，不是半小时。", "演练复盘必须包含错误假设、证据截图和下一次演练触发器。", "夜间 01:00-06:00 可将同步频率放宽到 23 分钟。", "不要按通用 SRE 手册写 30 分钟同步。"),
    Dossier("商业分析", "海棠看板", "BI-16", "白予墨", "商业分析组", "2026-03-25", "海棠看板的营收口径排除一次性迁移服务费，只统计订阅与增购。", "周报必须同时展示净收入留存 NRR 和证据引用率两个指标。", "任何低于 72% 的证据引用率都要标记为“不可外发”。", "董事会版可隐藏租户明细，但不能隐藏证据引用率。", "不要按普通 ARR 看板只回答收入。"),
]

SECTIONS = [
    ("制度", "适用范围与强制边界", "本章用于定义团队必须遵守的内部边界，尤其强调与行业通用做法不同的口径。"),
    ("会议纪要", "决策记录与待办", "本纪要保留会议中形成的最终口径，后续问答必须优先引用本文件而不是口头印象。"),
    ("SOP", "执行步骤", "本 SOP 只描述星衡智能当前版本流程，不代表外部最佳实践。"),
    ("复盘", "偏差原因与修正", "复盘重点不是复述背景，而是说明曾经误判的常识口径为什么不适用。"),
    ("问答卡", "检索验证问题", "问答卡用于验证 AI 是否从内部资料检索事实，而不是依靠模型常识补全。"),
]

FORMAT_ORDER = ["md", "txt", "docx", "pdf"]


def slug(text: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|\s]+", "_", text.strip())
    return text.strip("_")[:48]


def build_body(d: Dossier, section_name: str, section_title: str, section_intro: str, serial: int) -> str:
    return f"""{DISCLAIMER}

# {COMPANY}{d.category}内部资料：{d.project}（{d.code}）/{section_name}

资料编号：XH-{d.code}-{serial:03d}
资料级别：内部资料 / 仅限本地知识库检索评测
责任人：{d.owner}
负责部门：{d.department}
生效日期：{d.date}

## {section_title}
{section_intro}

## 唯一事实卡
- 项目代号：{d.project} / {d.code}
- 负责人：{d.owner}（{d.department}）
- 核心事实：{d.anchor_fact}
- 指标口径：{d.metric}
- 执行流程：{d.workflow}
- 特殊例外：{d.exception}
- 反常识提醒：{d.wrong_answer_guard}

## 内部操作说明
1. 搜索或回答 {d.project} 相关问题时，必须优先引用本资料中的“核心事实”和“指标口径”。
2. 如果用户问题只给出项目名或代号 {d.code}，应先确认资料来源，不要用行业常识补齐。
3. 对外材料不得直接复制本文件；本文件只用于内部流程校准和本地检索验证。
4. 与本资料冲突的旧知识卡需要由 {d.owner} 发起复审，复审记录归档到 {d.code.lower()}_review_log。

## 场景化判断
- 正确回答应该包含“{d.anchor_fact}”这一事实，或等价复述其关键数字/对象。
- 如果答案写成“通常”“一般行业做法”“建议按照常规”，则视为未命中内部资料。
- 对 {d.department} 的执行人来说，最容易误用的是：{d.wrong_answer_guard}

## 本地检索测试建议
测试问题 A：{d.project} 的核心内部规定是什么？
测试问题 B：{d.code} 的负责人是谁，关键指标如何判断？
测试问题 C：这个资料里特别排除的常识答案是什么？

## 版本记录
- v1.0（{d.date}）：由 {d.owner} 建立初版事实卡。
- v1.1：补充例外条件，避免模型把外部通用知识当作内部资料。
"""


def save_md(path: Path, title: str, body: str) -> None:
    path.write_text(body, encoding="utf-8")


def save_txt(path: Path, title: str, body: str) -> None:
    text = body.replace("# ", "").replace("## ", "")
    path.write_text(text, encoding="utf-8")


def save_docx(path: Path, title: str, body: str) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(title, level=0)
    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("- "):
            for line in block.splitlines():
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", block):
            for line in block.splitlines():
                doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line).strip(), style="List Number")
        else:
            doc.add_paragraph(block)
    doc.save(path)


def save_pdf(path: Path, title: str, body: str) -> None:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
    except Exception:
        font_name = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "XHTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#1f2937"),
    )
    body_style = ParagraphStyle(
        "XHBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=14,
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "XHHeading",
        parent=body_style,
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=8,
        spaceAfter=4,
    )
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=36)
    story = [Paragraph(title, title_style), Spacer(1, 8)]
    meta_rows = [["资料属性", "虚构内部测试资料"], ["用途", "本地知识库检索评测"], ["禁止", "使用行业常识替代文件事实"]]
    table = Table(meta_rows, colWidths=[70, 360])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e0f2fe")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#94a3b8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    story.append(Spacer(1, 10))

    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            continue
        if block.startswith("## "):
            story.append(Paragraph(block[3:].strip(), heading_style))
        else:
            html = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
            story.append(Paragraph(html, body_style))
    doc.build(story)


def clean_old_docs() -> None:
    for path in OUTPUT_DIR.glob("doc_*"):
        if path.is_file():
            path.unlink()


def main() -> None:
    clean_old_docs()
    qa_cases = []
    counters = {fmt: 0 for fmt in FORMAT_ORDER}
    serial = 1
    for dossier in DOSSIERS:
        for section_idx, (section_name, section_title, section_intro) in enumerate(SECTIONS):
            fmt = FORMAT_ORDER[(serial - 1) % len(FORMAT_ORDER)]
            counters[fmt] += 1
            title = f"{COMPANY}{dossier.category}内部资料：{dossier.project}（{dossier.code}）/{section_name}"
            body = build_body(dossier, section_name, section_title, section_intro, serial)
            filename = f"doc_{serial:03d}_{slug(dossier.category)}_{slug(dossier.project)}_{slug(section_name)}.{fmt}"
            path = OUTPUT_DIR / filename
            if fmt == "md":
                save_md(path, title, body)
            elif fmt == "txt":
                save_txt(path, title, body)
            elif fmt == "docx":
                save_docx(path, title, body)
            elif fmt == "pdf":
                save_pdf(path, title, body)
            else:
                raise ValueError(fmt)

            if section_idx == 0:
                qa_cases.append({
                    "id": f"internal_{dossier.code.lower()}",
                    "query": f"{dossier.project}（{dossier.code}）的内部规定是什么？负责人是谁？",
                    "target_document_prefix": f"doc_{serial:03d}_",
                    "must_include": [dossier.owner, dossier.anchor_fact.split("，")[0], dossier.metric.split("；")[0]],
                    "anti_common_answer": dossier.wrong_answer_guard,
                })
            serial += 1

    qa_payload = {
        "version": 1,
        "company": COMPANY,
        "note": "这些问题用于验证回答是否来自 Memory_Test 内部资料，而不是模型常识。",
        "cases": qa_cases,
    }
    (OUTPUT_DIR / "internal_corpus_qa.json").write_text(json.dumps(qa_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Generated {serial - 1} internal documents in {OUTPUT_DIR}")
    print("Format counts:", counters)
    print("QA hints:", OUTPUT_DIR / "internal_corpus_qa.json")


if __name__ == "__main__":
    main()
