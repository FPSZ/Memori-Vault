#!/usr/bin/env python3
"""
Generate a deterministic 5,000-file stress corpus for Memori-Vault.

Output layout:
  Memory_Test/stress_5000/
    md/      1,250 Markdown files
    txt/     1,250 text files
    docx/    1,250 Word files
    pdf/     1,250 PDF files
    manifest.json

The corpus is intentionally synthetic and local-first: it avoids downloading
large or copyrighted random internet documents while still exercising parser,
chunking, indexing, CJK, English, path, code-token, and mixed-format behavior.
"""

from __future__ import annotations

import argparse
import json
import random
import re
import shutil
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_OUTPUT_DIR = SCRIPT_DIR / "stress_5000"
FORMATS = ("md", "txt", "docx", "pdf")
FILES_PER_FORMAT = 1250
RANDOM_SEED = 20260425


@dataclass(frozen=True)
class Topic:
    category: str
    language: str
    title: str
    keywords: tuple[str, ...]


TOPICS: tuple[Topic, ...] = (
    Topic("ai", "zh", "人工智能", ("RAG", "embedding", "rerank", "agent", "模型治理")),
    Topic("ai", "en", "Artificial Intelligence", ("RAG", "embedding", "rerank", "agent", "alignment")),
    Topic("retrieval", "zh", "信息检索", ("BM25", "RRF", "向量召回", "证据链", "检索评测")),
    Topic("retrieval", "en", "Information Retrieval", ("BM25", "RRF", "dense retrieval", "citation", "evaluation")),
    Topic("database", "zh", "数据库", ("SQLite", "FTS5", "事务", "索引", "备份恢复")),
    Topic("database", "en", "Database Systems", ("SQLite", "FTS5", "transaction", "index", "backup")),
    Topic("security", "zh", "安全合规", ("RBAC", "audit", "egress policy", "最小权限", "密钥管理")),
    Topic("security", "en", "Security Compliance", ("RBAC", "audit", "egress policy", "least privilege", "key management")),
    Topic("product", "zh", "产品设计", ("onboarding", "可信度面板", "图谱可视化", "可解释性", "用户路径")),
    Topic("product", "en", "Product Design", ("onboarding", "trust panel", "graph view", "explainability", "journey")),
    Topic("architecture", "zh", "技术架构", ("Tauri", "Rust", "pipeline", "backpressure", "observability")),
    Topic("architecture", "en", "System Architecture", ("Tauri", "Rust", "pipeline", "backpressure", "observability")),
    Topic("operations", "zh", "运营增长", ("cohort", "retention", "activation", "north star metric", "实验")),
    Topic("operations", "en", "Growth Operations", ("cohort", "retention", "activation", "north star metric", "experiment")),
    Topic("project", "zh", "项目管理", ("sprint", "风险登记", "里程碑", "复盘", "验收标准")),
    Topic("project", "en", "Project Management", ("sprint", "risk register", "milestone", "retro", "acceptance criteria")),
    Topic("robotics", "zh", "机器人", ("SLAM", "传感器融合", "控制系统", "路径规划", "边缘计算")),
    Topic("robotics", "en", "Robotics", ("SLAM", "sensor fusion", "control loop", "path planning", "edge computing")),
    Topic("software", "zh", "软件工程", ("API", "CI/CD", "测试金字塔", "可维护性", "重构")),
    Topic("software", "en", "Software Engineering", ("API", "CI/CD", "test pyramid", "maintainability", "refactor")),
)


ZH_PARAGRAPHS = (
    "本文件用于压力测试本地知识库的解析、分块、向量化和证据链能力。内容包含中文、英文缩写、路径、API 名称与编号，方便发现检索排序中的边界问题。",
    "在混合语料场景中，系统需要先判断候选文档，再进入 chunk 级召回。若 document routing 不稳定，后续 rerank 和 citation 都会被错误候选拖偏。",
    "可靠的本地记忆系统应该在没有证据时拒答，在证据充分时给出可追溯引用，并把命中文档、命中原因、召回分数和过滤规则展示给用户。",
    "图谱可视化适合用于 evidence exploration：节点来自实体，边来自关系，来源必须能反查到 chunk。图谱不应在早期影响主召回排序，以免引入不可控噪声。",
    "索引系统需要处理暂停、恢复、重建、失败文件、过期任务和模型超时。任何后台队列都要有 backpressure 策略，避免通道写满导致死锁。",
    "企业私有化部署更关注审计、权限、备份恢复和模型外联策略。默认本地优先，所有敏感内容不应在未授权情况下发送到公网服务。",
)


EN_PARAGRAPHS = (
    "This document is part of a synthetic stress corpus for local-first memory retrieval. It includes English terms, CJK snippets, code-like tokens, and source paths for ranking diagnostics.",
    "A robust retrieval pipeline should separate document-level routing from chunk-level evidence selection. Weak routing often causes citation drift even when the generator is well prompted.",
    "Hybrid search combines lexical matching, dense embeddings, reciprocal rank fusion, and gating. Each stage should emit metrics so ranking failures can be classified rather than guessed.",
    "Graph visualization is useful for evidence exploration. Entities, relationships, source chunks, and document aggregation must remain explainable and reversible.",
    "Indexing pipelines need explicit progress, retry policy, timeout handling, and safe rebuild semantics. Long-running model calls should never make the application appear silently stuck.",
    "Agent integrations through MCP should expose ask, search, get_source, open_source, diagnostics, and settings tools while clearly warning about full-control operations.",
)


CODE_SNIPPETS = (
    "GET /api/indexing/status",
    "POST /api/ask",
    "memori://chunk/{chunk_id}",
    "memori://source/{path}",
    "docs/retrieval_regression_suite.json",
    "src/retrieval.rs::hybrid_search",
    "rank_settings(query=\"embedding timeout\")",
    "RRF(k=60), document_top_k=20, chunk_top_k=8",
)


def slugify(text: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "_", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:80] or "document"


def build_body(topic: Topic, ordinal: int, rng: random.Random) -> str:
    paragraphs = list(ZH_PARAGRAPHS if topic.language == "zh" else EN_PARAGRAPHS)
    rng.shuffle(paragraphs)
    selected = paragraphs[: rng.randint(3, 5)]
    keyword_line = (
        "关键词：" if topic.language == "zh" else "Keywords:"
    ) + " " + ", ".join(topic.keywords)
    code_line = (
        "诊断线索：" if topic.language == "zh" else "Diagnostic hints:"
    ) + " " + "; ".join(rng.sample(CODE_SNIPPETS, 3))
    source_line = (
        f"来源路径样例：Memory_Test/stress_5000/{topic.category}/{ordinal:05d}.md"
        if topic.language == "zh"
        else f"Source path example: Memory_Test/stress_5000/{topic.category}/{ordinal:05d}.md"
    )
    numbered_fact = (
        f"样本编号 {ordinal:05d} 的固定事实：{topic.title} 与 {topic.keywords[0]} 存在强相关。"
        if topic.language == "zh"
        else f"Fixed fact for sample {ordinal:05d}: {topic.title} is strongly related to {topic.keywords[0]}."
    )
    return "\n\n".join([keyword_line, code_line, source_line, numbered_fact, *selected])


def save_md(path: Path, title: str, body: str) -> None:
    path.write_text(f"# {title}\n\n{body}\n", encoding="utf-8")


def save_txt(path: Path, title: str, body: str) -> None:
    path.write_text(f"{title}\n{'=' * len(title)}\n\n{body}\n", encoding="utf-8")


def save_docx(path: Path, title: str, body: str) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in body.split("\n\n"):
        text = paragraph.strip()
        if text:
            doc.add_paragraph(text)
    doc.save(path)


def escape_pdf_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def save_pdf(path: Path, title: str, body: str) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph(f"<b>{escape_pdf_text(title)}</b>", styles["Title"]), Spacer(1, 10)]
    for paragraph in body.split("\n\n"):
        text = paragraph.strip()
        if text:
            story.append(Paragraph(escape_pdf_text(text), styles["BodyText"]))
            story.append(Spacer(1, 6))
    doc.build(story)


def generate(output_dir: Path, clean: bool) -> dict[str, object]:
    if clean and output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    rng = random.Random(RANDOM_SEED)
    manifest: dict[str, object] = {
        "seed": RANDOM_SEED,
        "files_per_format": FILES_PER_FORMAT,
        "total_files": FILES_PER_FORMAT * len(FORMATS),
        "formats": {},
        "note": "Synthetic local stress corpus for parser/index/retrieval performance tests.",
    }

    for fmt in FORMATS:
        fmt_dir = output_dir / fmt
        fmt_dir.mkdir(parents=True, exist_ok=True)
        entries = []
        for index in range(1, FILES_PER_FORMAT + 1):
            topic = rng.choice(TOPICS)
            title = f"{topic.title} stress sample {index:04d}"
            body = build_body(topic, index, rng)
            filename = f"{fmt}_{index:04d}_{topic.language}_{slugify(topic.title)}.{fmt}"
            path = fmt_dir / filename

            if fmt == "md":
                save_md(path, title, body)
            elif fmt == "txt":
                save_txt(path, title, body)
            elif fmt == "docx":
                save_docx(path, title, body)
            elif fmt == "pdf":
                pdf_topic = topic if topic.language == "en" else rng.choice([t for t in TOPICS if t.language == "en"])
                save_pdf(path, f"{pdf_topic.title} stress sample {index:04d}", build_body(pdf_topic, index, rng))
            else:
                raise ValueError(f"unsupported format: {fmt}")

            entries.append(
                {
                    "file": str(path.relative_to(output_dir)).replace("\\", "/"),
                    "category": topic.category,
                    "language": topic.language,
                    "title": title,
                }
            )
        manifest["formats"][fmt] = {
            "count": len(entries),
            "directory": str(fmt_dir.relative_to(output_dir)).replace("\\", "/"),
            "sample": entries[:5],
        }
        print(f"{fmt}: generated {len(entries)} files in {fmt_dir}")

    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def main() -> None:
    parser = argparse.ArgumentParser(
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description=textwrap.dedent(__doc__).strip(),
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--clean", action="store_true", help="remove the output directory before generating")
    args = parser.parse_args()

    manifest = generate(args.output, args.clean)
    print(json.dumps({"output": str(args.output), "total_files": manifest["total_files"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
